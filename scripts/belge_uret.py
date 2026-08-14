"""
TaskFlow teslim belgesini (.docx) uretir.

Yonergedeki uc temel cikti tek belgede toplanir:
  1) Proje Tanitim Dokumani
  2) Kurulum Notu (README)
  3) Postman test ekran goruntuleri
Ayrica kapakta GitHub adresi yer alir (form kod dosyasi kabul etmiyor).

Calistirma:
    python scripts/belge_uret.py

Cikti: OmerGunes_TaskFlow_Proje.docx  (Word'de acilip PDF'e cevrilecek)
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJE = Path(__file__).resolve().parent.parent
GORSELLER = PROJE / 'postman' / 'ekran_goruntuleri'
CIKTI = PROJE / 'OmerGunes_TaskFlow_Proje.docx'

GITHUB = 'https://github.com/MRGNSs11/taskflow-api'

# --- renkler ---
TURKUAZ = RGBColor(0x0E, 0x80, 0x8C)
KOYU = RGBColor(0x0F, 0x17, 0x2A)
SOLGUN = RGBColor(0x64, 0x74, 0x8B)


# ---------------------------------------------------------------- yardimcilar
def golge(hucre, renk_hex):
    """Tablo hucresine arka plan rengi verir."""
    dolgu = OxmlElement('w:shd')
    dolgu.set(qn('w:val'), 'clear')
    dolgu.set(qn('w:fill'), renk_hex)
    hucre._tc.get_or_add_tcPr().append(dolgu)


def kod_kutusu(belge, metin):
    """Kodu tek hucreli gri tabloda, Consolas ile gosterir."""
    tablo = belge.add_table(rows=1, cols=1)
    tablo.alignment = WD_TABLE_ALIGNMENT.CENTER
    hucre = tablo.cell(0, 0)
    golge(hucre, 'F4F6F8')

    hucre.text = ''
    ilk = True
    for satir in metin.strip('\n').split('\n'):
        p = hucre.paragraphs[0] if ilk else hucre.add_paragraph()
        ilk = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        calisma = p.add_run(satir)
        calisma.font.name = 'Consolas'
        calisma.font.size = Pt(9)
        calisma.font.color.rgb = KOYU

    belge.add_paragraph()


def tablo_ekle(belge, basliklar, satirlar, genislikler=None):
    """Basligi renkli, kenarli bir tablo ekler."""
    tablo = belge.add_table(rows=1, cols=len(basliklar))
    tablo.style = 'Table Grid'
    tablo.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, baslik in enumerate(basliklar):
        hucre = tablo.rows[0].cells[i]
        hucre.text = ''
        p = hucre.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        calisma = p.add_run(baslik)
        calisma.bold = True
        calisma.font.size = Pt(9.5)
        calisma.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        golge(hucre, '0E808C')

    for satir in satirlar:
        hucreler = tablo.add_row().cells
        for i, deger in enumerate(satir):
            hucreler[i].text = ''
            p = hucreler[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            calisma = p.add_run(str(deger))
            calisma.font.size = Pt(9.5)

    if genislikler:
        for satir in tablo.rows:
            for i, genislik in enumerate(genislikler):
                satir.cells[i].width = Cm(genislik)

    belge.add_paragraph()


def gorsel_ekle(belge, dosya_adi, resim_yazisi):
    """Ekran goruntusu ve altina italik resim yazisi ekler."""
    yol = GORSELLER / dosya_adi

    if not yol.exists():
        p = belge.add_paragraph()
        c = p.add_run(f'[EKSIK GORSEL: {dosya_adi}]')
        c.bold = True
        c.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
        return False

    belge.add_picture(str(yol), width=Cm(16.5))
    belge.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = belge.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = p.add_run(resim_yazisi)
    c.italic = True
    c.font.size = Pt(9)
    c.font.color.rgb = SOLGUN
    return True


def baslik1(belge, metin):
    p = belge.add_heading(metin, level=1)
    for c in p.runs:
        c.font.color.rgb = TURKUAZ
    return p


def baslik2(belge, metin):
    p = belge.add_heading(metin, level=2)
    for c in p.runs:
        c.font.color.rgb = KOYU
    return p


def paragraf(belge, metin, punto=10.5):
    p = belge.add_paragraph()
    c = p.add_run(metin)
    c.font.size = Pt(punto)
    p.paragraph_format.space_after = Pt(8)
    return p


def madde(belge, metin):
    p = belge.add_paragraph(style='List Bullet')
    c = p.add_run(metin)
    c.font.size = Pt(10.5)
    return p


def sayfa_sonu(belge):
    belge.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def sayfa_numarasi_ekle(belge):
    """Altbilgiye 'Sayfa X' alani koyar."""
    altbilgi = belge.sections[0].footer
    p = altbilgi.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    c = p.add_run('Sayfa ')
    c.font.size = Pt(9)
    c.font.color.rgb = SOLGUN

    alan = OxmlElement('w:fldSimple')
    alan.set(qn('w:instr'), 'PAGE')
    p._p.append(alan)


# ---------------------------------------------------------------------- belge
def main() -> None:
    belge = Document()

    # Sayfa kenar bosluklari
    for bolum in belge.sections:
        bolum.top_margin = Cm(2.2)
        bolum.bottom_margin = Cm(2.2)
        bolum.left_margin = Cm(2.2)
        bolum.right_margin = Cm(2.2)

    # Varsayilan yazi tipi
    stil = belge.styles['Normal']
    stil.font.name = 'Calibri'
    stil.font.size = Pt(10.5)

    sayfa_numarasi_ekle(belge)

    # ============================================================== KAPAK
    for _ in range(4):
        belge.add_paragraph()

    p = belge.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = p.add_run('TASKFLOW')
    c.bold = True
    c.font.size = Pt(38)
    c.font.color.rgb = TURKUAZ

    p = belge.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = p.add_run('Görev ve Proje Yönetim Sistemi API’si')
    c.font.size = Pt(16)
    c.font.color.rgb = KOYU

    p = belge.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = p.add_run('Node.js ile Backend Programlama — Bitirme Projesi')
    c.font.size = Pt(11.5)
    c.font.color.rgb = SOLGUN

    for _ in range(3):
        belge.add_paragraph()

    tablo_ekle(
        belge,
        ['', ''],
        [
            ['Hazırlayan', 'Ömer Güneş'],
            ['Teknolojiler', 'Node.js 24 · Express.js 5 · REST API'],
            ['Veri saklama', 'JSON dosyası (fs modülü)'],
            ['Tarih', 'Ağustos 2026'],
        ],
        genislikler=[5.0, 11.5],
    )

    belge.add_paragraph()

    p = belge.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = p.add_run('Kaynak Kod')
    c.bold = True
    c.font.size = Pt(12)
    c.font.color.rgb = KOYU

    p = belge.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = p.add_run(GITHUB)
    c.font.size = Pt(13)
    c.bold = True
    c.font.color.rgb = TURKUAZ

    p = belge.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = p.add_run('Depo geneldir (public); package.json ve proje klasörünün tamamı bu adrestedir.')
    c.font.size = Pt(9.5)
    c.font.color.rgb = SOLGUN

    sayfa_sonu(belge)

    # ======================================================== İÇİNDEKİLER
    baslik1(belge, 'İçindekiler')
    tablo_ekle(
        belge,
        ['Bölüm', 'Başlık'],
        [
            ['1', 'Proje Tanıtımı'],
            ['2', 'Veri Modeli'],
            ['3', 'API Tasarımı ve Mimari'],
            ['4', 'Logger Middleware (Zorunlu)'],
            ['5', 'Postman Testleri'],
            ['6', 'Kurulum Notu'],
            ['7', 'Klasör Yapısı'],
            ['8', 'İsteğe Bağlı Geliştirmeler'],
            ['—', 'Sonuç'],
        ],
        genislikler=[2.5, 14.0],
    )
    sayfa_sonu(belge)

    # ============================================================ BÖLÜM 1
    baslik1(belge, '1. Proje Tanıtımı')

    baslik2(belge, '1.1 Projenin Amacı')
    paragraf(
        belge,
        'Bu proje, eğitim boyunca öğrenilen Node.js, Express.js, REST API '
        'tasarımı, routing, middleware ve CRUD konularının tek bir uygulama '
        'üzerinde gösterilmesi amacıyla geliştirilmiştir.',
    )
    paragraf(
        belge,
        'Ortaya çıkan sistem, bir yazılım şirketinin ekip içi görev takibini '
        'yürütebileceği, tamamen REST API mantığında çalışan bir sunucu '
        'uygulamasıdır. Arayüzü yoktur; tüm işlemler HTTP istekleriyle yapılır '
        've JSON cevapları döner.',
    )

    baslik2(belge, '1.2 Senaryo')
    paragraf(
        belge,
        'Bir yazılım şirketi, ekip içerisindeki görevleri, projeleri ve '
        'çalışanların sorumluluklarını tek bir yerden takip etmek istemektedir. '
        'Hâlihazırda görevler sözlü olarak veya dağınık dosyalarda takip '
        'edilmekte; bu yüzden hangi işin kimde olduğu, hangisinin tamamlandığı '
        've hangisinin aciliyeti bulunduğu net görülememektedir.',
    )
    paragraf(belge, 'TaskFlow bu ihtiyacı karşılamak için tasarlanmıştır. Sistem üzerinden:')
    for m in [
        'Yeni görev oluşturulabilir',
        'Görevler çalışanlara atanabilir',
        'Görevlerin durumu takip edilebilir (beklemede, devam ediyor, tamamlandı, iptal edildi)',
        'Görevlere öncelik verilebilir (düşük, orta, yüksek, acil)',
        'Görevler duruma, önceliğe ve çalışana göre filtrelenebilir',
        'Temel raporlar alınabilir',
    ]:
        madde(belge, m)

    paragraf(
        belge,
        'Sistem, bir başka uygulamanın (web paneli, mobil uygulama) veri '
        'kaynağı olarak kullanılmak üzere tasarlanmıştır; bu nedenle yalnızca '
        'API katmanı geliştirilmiştir.',
    )

    baslik2(belge, '1.3 Sistem Özeti')
    tablo_ekle(
        belge,
        ['Başlık', 'Değer'],
        [
            ['Mimari', 'REST API, katmanlı yapı'],
            ['Çalışma ortamı', 'Node.js 18 ve üzeri'],
            ['Çerçeve', 'Express.js 5'],
            ['Veri saklama', 'JSON dosyası (src/data/gorevler.json)'],
            ['Kimlik doğrulama', 'Yok (yönerge kapsamı dışında)'],
            ['Bağımlılık sayısı', '1 (express)'],
        ],
        genislikler=[5.0, 11.5],
    )
    sayfa_sonu(belge)

    # ============================================================ BÖLÜM 2
    baslik1(belge, '2. Veri Modeli')
    paragraf(belge, 'Sistemde tek bir ana varlık vardır: Görev (Task).')

    tablo_ekle(
        belge,
        ['Alan', 'Tip', 'Zorunlu', 'Açıklama'],
        [
            ['id', 'number', '—', 'Otomatik artan, sistem üretir'],
            ['title', 'string', 'Evet', 'Görev başlığı (3–120 karakter)'],
            ['description', 'string', 'Hayır', 'Ayrıntılı açıklama (en fazla 1000 karakter)'],
            ['status', 'string', 'Hayır', 'pending · in-progress · completed · cancelled'],
            ['priority', 'string', 'Hayır', 'low · medium · high · urgent'],
            ['assignee', 'string', 'Hayır', 'Görevin atandığı çalışan (en fazla 60 karakter)'],
            ['createdAt', 'string', '—', 'Oluşturulma zamanı (ISO 8601)'],
            ['updatedAt', 'string', '—', 'Son güncelleme zamanı (ISO 8601)'],
        ],
        genislikler=[3.0, 2.2, 2.0, 9.3],
    )

    paragraf(
        belge,
        'Yeni görev oluşturulurken status verilmezse pending, priority '
        'verilmezse medium kabul edilir.',
    )

    baslik2(belge, 'Örnek Görev Kaydı')
    kod_kutusu(belge, '''{
  "id": 2,
  "title": "Veritabani semasinin olusturulmasi",
  "description": "Kullanici, gorev ve proje tablolarinin iliskileriyle birlikte kurulmasi.",
  "status": "in-progress",
  "priority": "urgent",
  "assignee": "Mehmet Kaya",
  "createdAt": "2026-08-05T11:00:00.000Z",
  "updatedAt": "2026-08-13T10:20:00.000Z"
}''')
    sayfa_sonu(belge)

    # ============================================================ BÖLÜM 3
    baslik1(belge, '3. API Tasarımı ve Mimari')

    baslik2(belge, '3.1 Katmanlı Yapı')
    paragraf(belge, 'Bir isteğin sistemde izlediği yol:')
    kod_kutusu(belge, 'İstek  →  Route  →  Middleware  →  Controller  →  Service  →  Data  →  Cevap')

    tablo_ekle(
        belge,
        ['Katman', 'Sorumluluk'],
        [
            ['Route', 'Hangi adresin hangi controller’a gideceği'],
            ['Middleware', 'İstek kaydı (logger) ve veri doğrulama'],
            ['Controller', 'İsteği alma, servisi çağırma, cevabı döndürme'],
            ['Service', 'İş kuralları: filtreleme, sıralama, güncelleme mantığı'],
            ['Data', 'JSON dosyasından okuma ve yazma'],
        ],
        genislikler=[3.5, 13.0],
    )

    paragraf(
        belge,
        'Bu ayrımın nedeni, her katmanın tek bir işten sorumlu olmasıdır. '
        'Örneğin ileride veri JSON dosyası yerine bir veritabanında tutulmak '
        'istenirse yalnızca Data katmanı değiştirilir; controller ve servis '
        'kodlarına dokunulmaz.',
    )

    baslik2(belge, '3.2 Zorunlu Uçlar')
    tablo_ekle(
        belge,
        ['Metot', 'Adres', 'Açıklama', 'Başarılı yanıt'],
        [
            ['POST', '/api/tasks', 'Görev ekleme', '201 Created'],
            ['GET', '/api/tasks', 'Görev listeleme', '200 OK'],
            ['GET', '/api/tasks/:id', 'Görev detayı', '200 OK'],
            ['PUT', '/api/tasks/:id', 'Görev güncelleme', '200 OK'],
            ['DELETE', '/api/tasks/:id', 'Görev silme', '204 No Content'],
        ],
        genislikler=[2.2, 4.3, 5.5, 4.5],
    )

    baslik2(belge, '3.3 Yanıt Zarfı')
    paragraf(
        belge,
        'Tüm uçlar aynı biçimi döndürür. Böylece API’yi kullanan istemci her '
        'ucu ayrı ayrı öğrenmek zorunda kalmaz; hata da başarı da aynı şekilde '
        'okunur.',
    )
    kod_kutusu(belge, '''{ "success": true,  "message": "...", "data": ... }

{ "success": false, "error": { "code": "...", "message": "...", "details": [...] } }''')

    baslik2(belge, '3.4 HTTP Durum Kodları')
    tablo_ekle(
        belge,
        ['Kod', 'Ne zaman döner'],
        [
            ['200', 'Başarılı okuma veya güncelleme'],
            ['201', 'Yeni görev oluşturuldu'],
            ['204', 'Silme başarılı, dönecek gövde yok'],
            ['400', 'Geçersiz veri veya geçersiz görev numarası'],
            ['404', 'Görev veya adres bulunamadı'],
            ['500', 'Beklenmeyen sunucu hatası'],
        ],
        genislikler=[2.5, 14.0],
    )

    baslik2(belge, '3.5 Hata Kodları')
    tablo_ekle(
        belge,
        ['error.code', 'HTTP', 'Ne zaman'],
        [
            ['VALIDATION_ERROR', '400', 'Alan kuralları sağlanmadı veya numara geçersiz'],
            ['INVALID_JSON', '400', 'İstek gövdesi geçerli JSON değil'],
            ['NOT_FOUND', '404', 'İstenen görev yok'],
            ['ROUTE_NOT_FOUND', '404', 'Böyle bir adres tanımlı değil'],
            ['INTERNAL_ERROR', '500', 'Beklenmeyen sunucu hatası'],
        ],
        genislikler=[4.5, 2.0, 10.0],
    )
    sayfa_sonu(belge)

    # ============================================================ BÖLÜM 4
    baslik1(belge, '4. Logger Middleware (Zorunlu)')
    paragraf(
        belge,
        'Yönergenin zorunlu kapsamındaki tek middleware budur. Sisteme gelen '
        'her isteğin HTTP metodu, adresi ve zaman damgası hem konsola hem '
        'logs/istekler.log dosyasına yazılır. Buna ek olarak dönen durum kodu '
        've isteğin kaç milisaniyede tamamlandığı da kaydedilmiştir.',
    )

    baslik2(belge, 'Örnek Kayıt Satırları')
    kod_kutusu(belge, '''[2026-08-14T11:09:02.034Z] GET /api/tasks/abc -> 400 (1ms)
[2026-08-14T11:09:02.074Z] GET /api/olmayan-uc -> 404 (1ms)
[2026-08-14T11:09:23.651Z] GET /api/tasks?status=completed -> 200 (0ms)
[2026-08-14T11:09:23.861Z] GET /api/reports/summary -> 200 (1ms)''')

    baslik2(belge, 'Nasıl Çalışır')
    paragraf(
        belge,
        'Middleware, Express’in res.on("finish") olayına bağlanır. Böylece '
        'kayıt, cevap gönderildikten sonra yazılır ve durum kodu ile geçen '
        'süre de bilinir. Log dosyasına yazma başarısız olursa istek yine de '
        'normal şekilde tamamlanır; kayıt tutulamaması kullanıcıyı etkilemez.',
    )

    baslik2(belge, 'Middleware Sırası')
    paragraf(belge, 'Sıra önemlidir; app.js içinde şu düzende bağlanmıştır:')
    kod_kutusu(belge, '''1) express.json()      Gelen JSON gövdesini çözümler
2) logger              ZORUNLU: her isteği kaydeder
3) rotalar             /api altındaki tüm uçlar
4) bulunamadi          Hiçbir rotaya uymayan istekler (404)
5) hataYakala          Merkezî hata yakalayıcı (en sonda olmalı)''')
    sayfa_sonu(belge)

    # ============================================================ BÖLÜM 5
    baslik1(belge, '5. Postman Testleri')
    paragraf(
        belge,
        'Yönergede istendiği gibi beş zorunlu ucun (ekleme, listeleme, detay, '
        'güncelleme, silme) tamamı Postman üzerinden test edilmiş ve doğru '
        'çalıştığı ekran görüntüleriyle belgelenmiştir. Ayrıca hata durumu da '
        'ayrı bir istekle doğrulanmıştır.',
    )
    paragraf(
        belge,
        'Hazır istek koleksiyonu depoda yer almaktadır: '
        'postman/TaskFlow_API.postman_collection.json',
    )

    testler = [
        ('postman-1.png', '5.1', 'Görev Ekleme — POST /api/tasks',
         'Yeni görev oluşturulur ve 201 Created döner. Cevapta sistemin ürettiği '
         'id, createdAt ve updatedAt alanları görülmektedir.',
         'Şekil 5.1 — POST /api/tasks isteği ve 201 Created cevabı'),
        ('postman-2.png', '5.2', 'Görev Listeleme — GET /api/tasks',
         'Sistemdeki tüm görevler listelenir ve 200 OK döner. Cevap zarfında '
         'kayıt sayısı da yer alır.',
         'Şekil 5.2 — GET /api/tasks isteği ve 200 OK cevabı'),
        ('postman-3.png', '5.3', 'Görev Detayı — GET /api/tasks/:id',
         'Belirli bir görev numarası ile tek kayıt getirilir ve 200 OK döner.',
         'Şekil 5.3 — GET /api/tasks/2 isteği ve 200 OK cevabı'),
        ('postman-4.png', '5.4', 'Görev Güncelleme — PUT /api/tasks/:id',
         'Yalnızca gönderilen alanlar güncellenir, updatedAt alanı yenilenir ve '
         '200 OK döner.',
         'Şekil 5.4 — PUT /api/tasks/2 isteği ve 200 OK cevabı'),
        ('postman-5.png', '5.5', 'Görev Silme — DELETE /api/tasks/:id',
         'Görev silinir ve 204 No Content döner. Bu durum kodunda cevap gövdesi '
         'bilinçli olarak boştur; standart budur.',
         'Şekil 5.5 — DELETE /api/tasks isteği ve 204 No Content cevabı'),
        ('postman-6.png', '5.6', 'Hata Durumu — Olmayan Görev (404)',
         'Var olmayan bir görev numarası istendiğinde sistem 404 Not Found ve '
         'açıklayıcı bir hata mesajı döndürür.',
         'Şekil 5.6 — GET /api/tasks/999 isteği ve 404 Not Found cevabı'),
    ]

    eksikler = []
    for dosya, no, baslik, aciklama, resim_yazisi in testler:
        baslik2(belge, f'{no} {baslik}')
        paragraf(belge, aciklama)
        if not gorsel_ekle(belge, dosya, resim_yazisi):
            eksikler.append(dosya)
        belge.add_paragraph()

    sayfa_sonu(belge)

    # ============================================================ BÖLÜM 6
    baslik1(belge, '6. Kurulum Notu')
    paragraf(belge, 'Gereksinimler: Node.js 18 veya üzeri ve npm.')

    baslik2(belge, '6.1 Kurulum')
    kod_kutusu(belge, f'''git clone {GITHUB}.git
cd taskflow-api
npm install''')
    paragraf(belge, 'Projenin tek bağımlılığı express’tir.')

    baslik2(belge, '6.2 Sunucuyu Çalıştırma')
    kod_kutusu(belge, 'npm start')
    paragraf(belge, 'Sunucu http://localhost:3000 adresinde çalışır. Açılışta konsola şunlar yazılır:')
    kod_kutusu(belge, '''  TaskFlow API
  Adres  : http://localhost:3000
  Uclar  : http://localhost:3000/api
  Kayit  : logs/istekler.log''')

    paragraf(belge, 'Farklı bir port kullanmak için:')
    kod_kutusu(belge, '''# Windows PowerShell
$env:PORT=4000; npm start

# macOS / Linux
PORT=4000 npm start''')
    paragraf(belge, 'Sunucuyu durdurmak için terminalde Ctrl + C tuşlanır.')

    baslik2(belge, '6.3 Çalıştığını Doğrulama')
    kod_kutusu(belge, 'curl http://localhost:3000/api')
    paragraf(belge, 'Uçların listesini içeren bir JSON cevabı dönerse API ayaktadır.')

    baslik2(belge, '6.4 Postman ile Test')
    for m in [
        'Postman’i açın ve Import düğmesine basın',
        'postman/TaskFlow_API.postman_collection.json dosyasını seçin',
        'Sunucunun çalıştığından emin olun (npm start)',
        'Koleksiyondaki istekleri sırayla çalıştırın',
    ]:
        madde(belge, m)
    sayfa_sonu(belge)

    # ============================================================ BÖLÜM 7
    baslik1(belge, '7. Klasör Yapısı')
    kod_kutusu(belge, '''src/
├── server.js                 Sunucuyu başlatan giriş noktası
├── app.js                    Express kurulumu ve middleware sırası
├── config/
│   └── ayarlar.js            Port, dosya yolları, sabit değerler
├── routes/                   Yalnızca adres tanımları
│   ├── index.js
│   ├── gorevRotalari.js
│   └── raporRotalari.js
├── controllers/              İstek alma ve cevap döndürme
│   ├── gorevController.js
│   └── raporController.js
├── services/
│   └── gorevServisi.js       İş mantığı
├── data/
│   ├── gorevDeposu.js        Veri erişimi
│   └── gorevler.json         Görevlerin saklandığı dosya
├── middleware/
│   ├── logger.js             ZORUNLU: istek kaydı
│   ├── dogrulama.js          İsteğe bağlı: alan doğrulama
│   └── hataYakala.js         Merkezî hata yakalama + 404
└── utils/
    ├── yanit.js              Tutarlı JSON yanıt zarfı
    └── ApiHatasi.js          Özel hata sınıfı

docs/                         Proje tanıtım ve API tasarım dokümanları
postman/                      Postman koleksiyonu ve ekran görüntüleri
logs/                         Çalışma sırasında üretilen istek kayıtları
package.json                  Proje tanımı ve bağımlılıklar''')
    sayfa_sonu(belge)

    # ============================================================ BÖLÜM 8
    baslik1(belge, '8. İsteğe Bağlı Geliştirmeler')

    p = belge.add_paragraph()
    c = p.add_run(
        'Aşağıdakiler yönergenin “İsteğe Bağlı” bölümünde yer alan, teslim ve '
        'değerlendirmeye dahil olmayan pratiklerdir. Zorunlu kapsam önceki '
        'bölümlerde tamamlanmıştır; bu bölüm ek olarak geliştirilmiştir.'
    )
    c.italic = True
    c.font.size = Pt(10)
    c.font.color.rgb = SOLGUN
    belge.add_paragraph()

    baslik2(belge, '8.1 Filtreleme, Arama, Sıralama, Sayfalama')
    paragraf(belge, 'Hepsi GET /api/tasks üzerinde sorgu parametresiyle çalışır ve birlikte kullanılabilir.')
    tablo_ekle(
        belge,
        ['Parametre', 'Örnek', 'Açıklama'],
        [
            ['status', '?status=completed', 'Duruma göre filtreleme'],
            ['priority', '?priority=high', 'Önceliğe göre filtreleme'],
            ['assignee', '?assignee=Omer Gunes', 'Çalışana göre listeleme'],
            ['keyword', '?keyword=api', 'Başlık ve açıklamada arama'],
            ['sort / order', '?sort=priority&order=desc', 'Sıralama'],
            ['page / limit', '?page=1&limit=10', 'Sayfalama (en fazla 100)'],
        ],
        genislikler=[3.2, 6.3, 7.0],
    )

    baslik2(belge, '8.2 Validation Middleware')
    paragraf(
        belge,
        'title, description, status, priority ve assignee alanları doğrulanır. '
        'Geçersiz veride 400 Bad Request ve alan bazlı açıklayıcı hata listesi döner.',
    )
    kod_kutusu(belge, '''{
  "success": false,
  "error": {
    "code": "VALIDATION_ERROR",
    "message": "Gonderilen veri gecerli degil.",
    "details": [ { "field": "title", "message": "Baslik zorunludur." } ]
  }
}''')

    baslik2(belge, '8.3 Raporlama Uçları')
    tablo_ekle(
        belge,
        ['Uç', 'Döndürdüğü'],
        [
            ['GET /api/reports/completed', 'Tamamlanan görev sayısı ve listesi'],
            ['GET /api/reports/pending', 'Bekleyen görev sayısı ve listesi'],
            ['GET /api/reports/summary', 'Duruma, önceliğe ve çalışana göre dağılım'],
        ],
        genislikler=[6.5, 10.0],
    )
    kod_kutusu(belge, '''{
  "toplamGorev": 6,
  "durumaGore": { "pending": 2, "in-progress": 2, "completed": 1, "cancelled": 1 },
  "oncelikeGore": { "low": 1, "medium": 2, "high": 2, "urgent": 1 },
  "calisanaGore": { "Ayse Demir": 2, "Mehmet Kaya": 2, "Omer Gunes": 2 }
}''')
    sayfa_sonu(belge)

    # =============================================================== SONUÇ
    baslik1(belge, 'Sonuç')
    paragraf(belge, 'Yönergede istenen dört zorunlu aşama da tamamlanmıştır.')
    tablo_ekle(
        belge,
        ['Aşama', 'Çıktı', 'Durum'],
        [
            ['1 — İhtiyaç analizi ve sistem tasarımı',
             'Proje tanıtım dokümanı, veri modeli, uç listesi', 'Tamamlandı'],
            ['2 — Backend altyapısı',
             'Çalışan Express projesi, katmanlı klasör yapısı, routing', 'Tamamlandı'],
            ['3 — CRUD + temel middleware',
             'Beş uç ve Logger middleware', 'Tamamlandı'],
            ['4 — Test ve teslim',
             'Postman testleri, kurulum notu, ekran görüntüleri', 'Tamamlandı'],
        ],
        genislikler=[5.5, 7.5, 3.5],
    )

    baslik2(belge, 'Teslim Edilen Çıktıların Bu Belgedeki Yeri')
    tablo_ekle(
        belge,
        ['İstenen çıktı', 'Bu belgede'],
        [
            ['Proje Tanıtım Dokümanı', 'Bölüm 1 ve Bölüm 2'],
            ['Kurulum Notu (README)', 'Bölüm 6'],
            ['Postman Test Ekran Görüntüleri', 'Bölüm 5'],
            ['Kaynak kodların tamamı', 'Kapaktaki GitHub adresi'],
        ],
        genislikler=[7.0, 9.5],
    )

    belge.save(CIKTI)

    print(f'Belge uretildi: {CIKTI.name}')
    print(f'Konum: {CIKTI}')
    if eksikler:
        print('\nEKSIK GORSELLER (belgede kirmizi uyari olarak isaretlendi):')
        for e in eksikler:
            print('  -', e)
    else:
        print('Tum ekran goruntuleri yerlestirildi.')


if __name__ == '__main__':
    main()
